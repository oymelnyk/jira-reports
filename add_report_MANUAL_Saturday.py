from docx import Document
#import os
from jira import JIRA
#from datetime import datetime, date, timedelta
#from pathlib import Path
#import shutil
#import time
from getpass import getpass
from docx.shared import Cm



def get_jira_info():
    #print("login: ")
    login = input()
    pwd = getpass()
    jira_server = ''
    jira= JIRA(server=jira_server, basic_auth=(login, pwd))

    assigned_groups = ['jira_it_GroupName_2line', 'jira_it_GroupName_2line', 'jira_it_GroupName_2line', 'jira_it_GroupName_2line', 'jira_it_GroupName_2line', 'jira_it_GroupName_2line', 'jira_it_GroupName_2line', 'jira_it_GroupName_2line', 'jira_it_GroupName_2line']
    
    global done_issues
    global in_progress_issues
    global canceled_issues


    done_issues = []
    in_progress_issues = []
    canceled_issues = []


    #jql_done = 'project = ITSD AND status in (Resolved, "Customer Approval") AND updated >= -7d AND updated <= 0d AND "Assigned group" = '
    #jql_in_progress = 'project = ITSD AND status in (Open, Reopened, "Waiting for customer", "Waiting for approval", "Work in progress", "Work in outsource", "Open 2 line") AND updated >= -7d AND updated <= 0d AND "Assigned group" = '
    #jql_canceled = 'project = ITSD AND status in (Closed, Canceled) AND updated >= -7d AND updated <= 0d AND "Assigned group" = '

    for assigned_group in assigned_groups:
        jql_done = jira.search_issues('project = ITSD AND status in (Resolved, "Customer Approval") AND updated >= -8d AND updated <= 0d AND "Assigned group" = '+ str(assigned_group), maxResults=10000)
        jql_in_progress = jira.search_issues('project = ITSD AND status in (Open, Reopened, "Waiting for customer", "Waiting for approval", "Work in progress", "Work in outsource", "Open 2 line") AND "Assigned group" = '+ str(assigned_group), maxResults=10000)
        jql_canceled = jira.search_issues('project = ITSD AND status in (Closed, Canceled) AND updated >= -8d AND updated <= 0d AND "Assigned group" = '+ str(assigned_group), maxResults=10000)
        #done_issues.append(jql_done)
        #in_progress_issues.append(len(jql_in_progress))
        #canceled_issues.append(len(jql_canceled))
        for issue in jql_done:
            done_issues.append(issue)
        for issue in jql_in_progress:
            in_progress_issues.append(issue)
        for issue in jql_canceled:
            canceled_issues.append(issue)
    
    
    #print("done")
    print(done_issues)
    #print(in_progress_issues)
    #print(canceled_issues)
    return done_issues, in_progress_issues, canceled_issues



def additionaly_file(done_issues, in_progress_issues, canceled_issues, save_path):
    d_i_len = len(done_issues)+1
    i_p_len = len(in_progress_issues)+1
    c_i_len = len(canceled_issues)+1
    document = Document()
    sections = document.sections

    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)

    #Table DONE
    
    def table_create(rows_len, cols_num, issues, table_name): #d_i_len,7
        p1 = document.add_paragraph()
        p1.add_run(str(table_name)).bold = True 
    
        table = document.add_table(rows=rows_len,cols=cols_num)
        table.style = 'TableGrid'

        header = ['Номер','Краткое Описание','Пользователь','Группа назначения','Исполнитель', 'Дата создания', 'Дата закрытия']
        hdr_cells = table.rows[0].cells
        for i in range(0,7):
            hdr_cells[i].text = str(header[i])

        ii=0
        for i in issues:
            print(str(i))
            ii=ii+1
            hdr_cells = table.rows[ii].cells
            hdr_cells[0].text = str(i)
            hdr_cells[1].text = str(i.fields.summary)
            hdr_cells[2].text = str(i.fields.reporter)
            hdr_cells[3].text = str(i.fields.customfield_10201)
            hdr_cells[4].text = str(i.fields.assignee)
            hdr_cells[5].text = str(i.fields.created).replace("T"," ").replace(".000+0300","")
            #hdr_cells[6].text = str(i.fields.resolved)
            try:
                hdr_cells[6].text = str(i.fields.resolutiondate).replace("T"," ").replace(".000+0300","")
            except:        
                hdr_cells[6].text = str("-")


    


    table_create(d_i_len, 7, done_issues, "Выполнено: ")
    table_create(i_p_len, 7, in_progress_issues, "В работе: " )
    table_create(c_i_len, 7, canceled_issues, "Отклонено: ")


    document.save(save_path)
get_jira_info()
additionaly_file(done_issues, in_progress_issues, canceled_issues, 'P:\\2021\\август\\Подробный отчет по заявкам 23-08_28-08-2021.docx')