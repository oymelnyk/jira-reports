from docx import Document
from datetime import datetime
from io import StringIO



def first():
    document = Document('C:\\Users\\Alexey\\Desktop\\new\\test.docx')
    p = document.add_paragraph(str(datetime.now()))

    records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam'))

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc




    document.save('C:\\Users\\Alexey\\Desktop\\new\\test1.docx')


#first()
def tst():
    import win32com.client as win32

    word = win32.gencache.EnsureDispatch('Word.Application')
    my_doc1=word.Documents.Open('C:\\Users\\Alexey\\Desktop\\new\\test1.docx')
    my_doc1.Visible = 0
    #my_doc=word.Documents.Open('C:\\Users\\Alexey\\Desktop\\new\\test.docx')
    #my_doc.Visible=True
    tables = my_doc1.Tables.Count
    print(tables)
    for table in range(1,tables+1):
    #    print(my_doc1.Tables(table).Cell(1,1).Range.Text)
    #    print(my_doc1.Tables(table).Cell(1,1).Range.Text)

    #print(my_doc1.Tables(2).Cell(1,1).Range.Text)
    #    print(table)

        print(str(my_doc1.Tables(table).Cell(1,1).Range.Text))
        if str('Группа') in str(my_doc1.Tables(table).Cell(1,1).Range.Text):
            print("+")
            my_doc = word.Documents.Open('C:\\Users\\Alexey\\Desktop\\new\\test.docx')
            my_doc.Visible = 0
            my_doc.Tables(1).Cell(1,1).Range.Text = my_doc1.Tables(table).Cell(1,1).Range.Text
            my_doc.Close()
            my_doc.Quit()
            my_doc1.Close()
            my_doc1.Quit()
        else:
            print("nf")

    


tst()
